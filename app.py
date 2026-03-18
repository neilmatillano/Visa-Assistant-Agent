"""
Visa Application Multi-Agent System — Phase 1
Streamlit Application (Dark Classy Theme · Tabular Step Layout)
"""

import io
import os
import sys

import streamlit as st

sys.path.insert(0, os.path.dirname(__file__))

from agents.orchestrator import VisaOrchestrator, PipelineStage
from agents.analysis_agent import VisaReport
from agents.tracker import generate_ics, build_tracker_pdf, STATUS_ICONS, STATUS_LABELS, ItemStatus, ItemType
import database as db


# ─────────────────────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Visa Assistant · AI-Powered",
    page_icon="🛂",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS — Dark classy theme + custom tab styling
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Playfair+Display:wght@600;700&display=swap');

:root {
    --bg-base:      #0d1117;
    --bg-surface:   #161b22;
    --bg-raised:    #1c2333;
    --bg-hover:     #21262d;
    --border:       #30363d;
    --border-light: #21262d;
    --gold:         #d4a843;
    --gold-light:   #f0c060;
    --gold-dim:     #8a6a20;
    --text-primary:   #e6edf3;
    --text-secondary: #8b949e;
    --text-muted:     #6e7681;
    --accent:    #58a6ff;
    --success:   #3fb950;
    --warning:   #d29922;
    --danger:    #f85149;
    --radius-sm: 8px;
    --radius-md: 12px;
    --radius-lg: 18px;
}

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif !important;
    color: var(--text-primary) !important;
}
.stApp { background-color: var(--bg-base) !important; }
.block-container { padding-top: 1.5rem !important; padding-bottom: 3rem !important; max-width: 1100px !important; }
#MainMenu, footer, header { visibility: hidden; }
hr { border: none !important; border-top: 1px solid var(--border) !important; margin: 20px 0 !important; }

/* ── Sidebar ── */
[data-testid="stSidebar"] { background-color: var(--bg-surface) !important; border-right: 1px solid var(--border) !important; }
[data-testid="stSidebar"] * { color: var(--text-primary) !important; }
[data-testid="stSidebar"] label { color: var(--text-secondary) !important; font-size: 0.75rem !important; font-weight: 600 !important; letter-spacing: 0.07em !important; text-transform: uppercase !important; }
[data-testid="stSidebar"] .stSelectbox > div > div { background-color: var(--bg-raised) !important; border: 1px solid var(--border) !important; color: var(--text-primary) !important; border-radius: var(--radius-sm) !important; }
[data-testid="stSidebar"] input { background-color: var(--bg-raised) !important; border: 1px solid var(--border) !important; color: var(--text-primary) !important; border-radius: var(--radius-sm) !important; }

/* ── Native Streamlit tabs — override to match dark theme ── */
.stTabs [data-baseweb="tab-list"] {
    background: var(--bg-surface) !important;
    border-radius: var(--radius-lg) var(--radius-lg) 0 0 !important;
    border: 1px solid var(--border) !important;
    border-bottom: none !important;
    padding: 0 8px !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    border: none !important;
    border-bottom: 3px solid transparent !important;
    border-radius: 0 !important;
    color: var(--text-secondary) !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    padding: 14px 22px !important;
    transition: all 0.2s !important;
    white-space: nowrap !important;
}
.stTabs [data-baseweb="tab"]:hover {
    color: var(--text-primary) !important;
    background: var(--bg-hover) !important;
}
.stTabs [aria-selected="true"] {
    color: var(--gold) !important;
    border-bottom: 3px solid var(--gold) !important;
    background: transparent !important;
}
.stTabs [data-baseweb="tab-panel"] {
    background: var(--bg-surface) !important;
    border: 1px solid var(--border) !important;
    border-top: none !important;
    border-radius: 0 0 var(--radius-lg) var(--radius-lg) !important;
    padding: 28px 32px !important;
}
/* Tab highlight bar override */
.stTabs [data-baseweb="tab-highlight"] { display: none !important; }
.stTabs [data-baseweb="tab-border"]    { display: none !important; }

/* ── Inputs ── */
.stTextInput input {
    background-color: var(--bg-raised) !important;
    border: 1px solid var(--border) !important;
    color: var(--text-primary) !important;
    border-radius: var(--radius-sm) !important;
    font-size: 0.93rem !important;
    padding: 10px 14px !important;
}
.stTextInput input:focus { border-color: var(--gold) !important; box-shadow: 0 0 0 3px rgba(212,168,67,0.15) !important; }
.stTextInput input::placeholder { color: var(--text-muted) !important; }
.stTextInput label { color: var(--text-secondary) !important; }
.stSelectbox > div > div { background-color: var(--bg-raised) !important; border: 1px solid var(--border) !important; color: var(--text-primary) !important; border-radius: var(--radius-sm) !important; }
.stSelectbox label { color: var(--text-secondary) !important; }

/* ── Buttons ── */
.stButton button { background-color: var(--bg-raised) !important; border: 1px solid var(--border) !important; color: var(--text-primary) !important; border-radius: var(--radius-sm) !important; font-weight: 600 !important; font-size: 0.88rem !important; transition: all 0.2s !important; }
.stButton button:hover { border-color: var(--gold) !important; color: var(--gold) !important; }
.stFormSubmitButton button { background: linear-gradient(135deg, var(--gold-dim), var(--gold)) !important; color: #0d1117 !important; border: none !important; border-radius: var(--radius-sm) !important; font-weight: 700 !important; padding: 10px 28px !important; box-shadow: 0 4px 14px rgba(212,168,67,0.3) !important; }
.stFormSubmitButton button:hover { box-shadow: 0 6px 20px rgba(212,168,67,0.45) !important; transform: translateY(-1px) !important; }
.stDownloadButton button { background: linear-gradient(135deg, #1a3a2a, var(--success)) !important; color: #fff !important; border: none !important; border-radius: var(--radius-sm) !important; font-weight: 700 !important; box-shadow: 0 4px 14px rgba(63,185,80,0.25) !important; }

/* ── Expander ── */
.streamlit-expanderHeader { background-color: var(--bg-raised) !important; border: 1px solid var(--border) !important; border-radius: var(--radius-sm) !important; color: var(--text-primary) !important; font-weight: 600 !important; font-size: 0.9rem !important; padding: 12px 16px !important; }
.streamlit-expanderHeader p, .streamlit-expanderHeader span, .streamlit-expanderHeader div { color: var(--text-primary) !important; }
.streamlit-expanderHeader svg { fill: var(--text-secondary) !important; }
.streamlit-expanderContent { background-color: var(--bg-surface) !important; border: 1px solid var(--border) !important; border-top: none !important; border-radius: 0 0 var(--radius-sm) var(--radius-sm) !important; padding: 14px 16px !important; }
.streamlit-expanderContent * { color: var(--text-secondary) !important; }

/* ── Spinner ── */
.stSpinner > div { border-top-color: var(--gold) !important; }

/* ── Markdown ── */
.stMarkdown p, .stMarkdown li, .stMarkdown span { color: var(--text-primary) !important; }

/* ─── CUSTOM COMPONENTS ─── */

.page-hero {
    background: linear-gradient(135deg, #1a1f2e 0%, #1c2333 100%);
    border: 1px solid var(--border);
    border-top: 3px solid var(--gold);
    border-radius: var(--radius-lg);
    padding: 28px 36px;
    margin-bottom: 24px;
    position: relative; overflow: hidden;
}
.page-hero::after { content:''; position:absolute; top:-60px; right:-60px; width:200px; height:200px; background:radial-gradient(circle,rgba(212,168,67,.08) 0%,transparent 70%); pointer-events:none; }
.page-hero h1 { margin:0 0 6px 0 !important; font-family:'Playfair Display',serif !important; font-size:1.9rem !important; font-weight:700 !important; color:var(--text-primary) !important; letter-spacing:-0.01em !important; }
.page-hero p  { margin:0 !important; color:var(--text-secondary) !important; font-size:0.9rem !important; }
.gold-line { color:var(--gold); font-weight:600; }

.section-label { font-size:0.7rem; font-weight:700; letter-spacing:0.1em; text-transform:uppercase; color:var(--text-muted); margin-bottom:12px; }

/* Step header inside tab */
.step-header {
    display: flex; align-items: center; gap: 14px;
    margin-bottom: 24px;
    padding-bottom: 18px;
    border-bottom: 1px solid var(--border);
}
.step-num-circle {
    width: 42px; height: 42px; border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.1rem; font-weight: 800; flex-shrink: 0;
}
.step-num-active  { background: linear-gradient(135deg,var(--gold-dim),var(--gold)); color:#0d1117; box-shadow:0 4px 14px rgba(212,168,67,.3); }
.step-num-done    { background: rgba(63,185,80,.15); color:var(--success); border:1px solid #1a5c2a; }
.step-num-waiting { background: var(--bg-hover); color:var(--text-muted); border:1px solid var(--border); }
.step-title { font-family:'Playfair Display',serif; font-size:1.25rem; font-weight:700; color:var(--text-primary); margin:0; }
.step-subtitle { font-size:0.82rem; color:var(--text-muted); margin:2px 0 0; }

/* Status pill */
.status-pill { display:inline-flex; align-items:center; gap:5px; padding:4px 12px; border-radius:20px; font-size:0.75rem; font-weight:600; margin-left:auto; }
.pill-active  { background:rgba(212,168,67,.12); color:var(--gold); border:1px solid var(--gold-dim); }
.pill-done    { background:rgba(63,185,80,.1); color:var(--success); border:1px solid #1a5c2a; }
.pill-waiting { background:var(--bg-hover); color:var(--text-muted); border:1px solid var(--border); }

/* Chat */
.chat-wrap { display:flex; flex-direction:column; gap:14px; padding:4px 0 12px; }
.msg-row-user  { display:flex; justify-content:flex-end; align-items:flex-end; gap:8px; }
.msg-row-agent { display:flex; justify-content:flex-start; align-items:flex-end; gap:8px; }
.msg-avatar { width:32px; height:32px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:.9rem; flex-shrink:0; border:1px solid var(--border); }
.avatar-agent { background:linear-gradient(135deg,#1a1f2e,#2a2f3e); }
.avatar-user  { background:linear-gradient(135deg,var(--gold-dim),var(--gold)); }
.msg-bubble { max-width:72%; padding:12px 16px; border-radius:14px; font-size:.9rem; line-height:1.6; }
.bubble-agent { background:var(--bg-raised); color:var(--text-primary) !important; border:1px solid var(--border); border-bottom-left-radius:4px; }
.bubble-user  { background:linear-gradient(135deg,var(--gold-dim),var(--gold)); color:#0d1117 !important; border-bottom-right-radius:4px; font-weight:500; }

/* Typing indicator */
.typing-bubble { display:flex !important; align-items:center; gap:6px; padding:14px 18px !important; min-width:64px; }
.typing-dot { display:inline-block; width:9px; height:9px; border-radius:50%; background:var(--gold); animation:typing-bounce 1.3s ease-in-out infinite; }
.typing-dot:nth-child(1) { animation-delay:0s; }
.typing-dot:nth-child(2) { animation-delay:0.22s; }
.typing-dot:nth-child(3) { animation-delay:0.44s; }
@keyframes typing-bounce { 0%,60%,100% { transform:translateY(0); opacity:0.4; } 30% { transform:translateY(-8px); opacity:1; } }

/* Research status card */
.research-card {
    background: var(--bg-raised);
    border: 1px solid var(--border);
    border-radius: var(--radius-md);
    padding: 24px 28px;
    text-align: center;
}
.research-card .rc-icon { font-size: 3rem; margin-bottom: 12px; }
.research-card .rc-title { font-family:'Playfair Display',serif; font-size:1.2rem; font-weight:700; color:var(--text-primary); margin-bottom:6px; }
.research-card .rc-sub   { font-size:.85rem; color:var(--text-secondary); }

/* Report hero */
.report-hero { background:linear-gradient(135deg,#1a1f2e 0%,#1c2333 100%); border:1px solid var(--border); border-left:4px solid var(--gold); border-radius:var(--radius-lg); padding:24px 28px; margin-bottom:20px; }
.report-hero h2 { margin:0 0 8px 0 !important; font-family:'Playfair Display',serif !important; font-size:1.5rem !important; font-weight:700 !important; color:var(--text-primary) !important; }
.report-hero p  { margin:0 !important; color:var(--text-secondary) !important; font-size:.87rem !important; }

/* Metric cards */
.metric-card { background:var(--bg-raised); border:1px solid var(--border); border-radius:var(--radius-md); padding:18px 20px; text-align:center; }
.metric-label { font-size:.68rem; font-weight:700; letter-spacing:.09em; text-transform:uppercase; color:var(--text-muted); margin-bottom:8px; }
.metric-value { font-size:1.05rem; font-weight:700; color:var(--gold) !important; }

/* Info banners */
.info-banner { border-radius:var(--radius-sm); padding:13px 16px; margin:10px 0; font-size:.87rem; line-height:1.55; }
.banner-blue  { background:#0d1f3c; border:1px solid #1d4a8a; color:#93c5fd !important; }
.banner-amber { background:#1f1500; border:1px solid #7a5000; color:#fbbf24 !important; }
.banner-green { background:#0a1f10; border:1px solid #1a5c2a; color:#6ee7b7 !important; }
.banner-red   { background:#1f0a0a; border:1px solid #7a1a1a; color:#fca5a5 !important; }

/* Steps */
.step-row { display:flex; gap:14px; padding:10px 0; align-items:flex-start; border-bottom:1px solid var(--border-light); }
.step-num  { background:linear-gradient(135deg,var(--gold-dim),var(--gold)); color:#0d1117; width:28px; height:28px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:.78rem; font-weight:800; flex-shrink:0; }
.step-text { font-size:.88rem; color:var(--text-primary) !important; line-height:1.55; padding-top:4px; }
.doc-detail { font-size:.84rem; color:var(--text-secondary) !important; line-height:1.5; }

/* Source links */
.source-link { display:inline-flex; align-items:center; gap:5px; background:var(--bg-raised); border:1px solid var(--border); border-radius:6px; padding:6px 12px; font-size:.81rem; font-weight:500; color:var(--accent) !important; text-decoration:none; margin:3px 3px 3px 0; transition:all .15s; }
.source-link:hover { border-color:var(--accent); background:#0d1f3c; }

/* ── Agent Activity Log Panel ── */
.log-panel {
    background: var(--bg-base);
    border: 1px solid var(--border);
    border-radius: var(--radius-md);
    padding: 0;
    overflow: hidden;
    margin-top: 20px;
}
.log-panel-header {
    background: var(--bg-raised);
    border-bottom: 1px solid var(--border);
    padding: 12px 18px;
    display: flex;
    align-items: center;
    gap: 10px;
}
.log-panel-title {
    font-size: .78rem;
    font-weight: 700;
    letter-spacing: .09em;
    text-transform: uppercase;
    color: var(--text-muted);
}
.log-count-badge {
    background: var(--bg-hover);
    border: 1px solid var(--border);
    border-radius: 10px;
    padding: 2px 8px;
    font-size: .72rem;
    font-weight: 600;
    color: var(--text-secondary);
    margin-left: auto;
}
.log-body {
    padding: 8px 0;
    max-height: 420px;
    overflow-y: auto;
}
.log-entry {
    display: flex;
    gap: 12px;
    padding: 9px 18px;
    border-bottom: 1px solid var(--border-light);
    align-items: flex-start;
    transition: background .1s;
}
.log-entry:last-child { border-bottom: none; }
.log-entry:hover { background: var(--bg-raised); }
.log-icon { font-size: 1rem; flex-shrink: 0; margin-top: 1px; }
.log-content { flex: 1; min-width: 0; }
.log-msg {
    font-size: .86rem;
    font-weight: 500;
    line-height: 1.45;
    color: var(--text-primary) !important;
    margin: 0;
}
.log-detail {
    font-size: .78rem;
    color: var(--text-secondary) !important;
    margin: 4px 0 0;
    line-height: 1.5;
    white-space: pre-wrap;
    font-family: 'Inter', monospace;
    background: var(--bg-raised);
    border: 1px solid var(--border-light);
    border-radius: 5px;
    padding: 6px 10px;
}
.log-ts {
    font-size: .7rem;
    color: var(--text-muted) !important;
    flex-shrink: 0;
    margin-top: 2px;
    font-family: monospace;
}
/* Level colour accents on left border */
.log-info    { border-left: 3px solid #30363d; }
.log-step    { border-left: 3px solid var(--gold-dim); }
.log-success { border-left: 3px solid var(--success); }
.log-warning { border-left: 3px solid var(--warning); }
.log-error   { border-left: 3px solid var(--danger); }

/* Sidebar */
.sidebar-logo { background:linear-gradient(135deg,#1a1f2e,#1c2333); border:1px solid var(--border); border-top:3px solid var(--gold); border-radius:var(--radius-md); padding:20px 16px; text-align:center; margin-bottom:20px; }
.sidebar-logo .logo-icon { font-size:2.4rem; }
.sidebar-logo .logo-title { font-family:'Playfair Display',serif !important; color:var(--text-primary) !important; font-weight:700 !important; font-size:1.15rem !important; margin-top:8px !important; }
.sidebar-logo .logo-sub { color:var(--text-muted) !important; font-size:.75rem !important; }
.visa-pill { display:inline-flex; align-items:center; gap:6px; background:var(--bg-raised); border:1px solid var(--border); border-radius:20px; padding:5px 12px; font-size:.82rem; font-weight:500; color:var(--text-secondary) !important; margin:3px 0; }
.disclaimer { font-size:.74rem; color:var(--text-muted) !important; line-height:1.5; }

/* ─────────────────────────────────────────────────────────────────────────────
   ANIMATIONS
───────────────────────────────────────────────────────────────────────────── */

/* Keyframes */
@keyframes pulse-ring {
    0%   { transform: scale(1);   opacity: 1; }
    60%  { transform: scale(1.55); opacity: 0; }
    100% { transform: scale(1.55); opacity: 0; }
}
@keyframes flow-right {
    0%   { transform: translateX(-100%); opacity: 0; }
    30%  { opacity: 1; }
    70%  { opacity: 1; }
    100% { transform: translateX(100%);  opacity: 0; }
}
@keyframes shimmer {
    0%   { background-position: -600px 0; }
    100% { background-position:  600px 0; }
}
@keyframes bounce-dots {
    0%, 80%, 100% { transform: translateY(0);    opacity: .4; }
    40%            { transform: translateY(-8px); opacity: 1; }
}
@keyframes spin-ring {
    to { transform: rotate(360deg); }
}
@keyframes fade-in-up {
    from { opacity: 0; transform: translateY(12px); }
    to   { opacity: 1; transform: translateY(0); }
}
@keyframes glow-pulse {
    0%, 100% { box-shadow: 0 0 0   0   rgba(212,168,67,.0); }
    50%       { box-shadow: 0 0 18px 4px rgba(212,168,67,.35); }
}

/* ── Pipeline connector bar ── */
.pipeline-bar {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 0;
    margin: 0 0 20px;
    padding: 18px 24px;
    background: var(--bg-surface);
    border: 1px solid var(--border);
    border-radius: var(--radius-lg);
    position: relative;
    overflow: hidden;
}
/* Shimmer overlay when any agent is active */
.pipeline-bar.running::before {
    content: '';
    position: absolute;
    inset: 0;
    background: linear-gradient(90deg,
        transparent 0%,
        rgba(212,168,67,.06) 40%,
        rgba(212,168,67,.12) 50%,
        rgba(212,168,67,.06) 60%,
        transparent 100%);
    background-size: 600px 100%;
    animation: shimmer 2.2s linear infinite;
    pointer-events: none;
    border-radius: inherit;
}

/* Individual pipeline node */
.pnode {
    display: flex;
    flex-direction: column;
    align-items: center;
    gap: 8px;
    min-width: 110px;
    position: relative;
    z-index: 1;
    animation: fade-in-up .5s ease both;
}
.pnode-icon-wrap {
    position: relative;
    width: 56px; height: 56px;
    display: flex; align-items: center; justify-content: center;
}
.pnode-circle {
    width: 56px; height: 56px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.4rem;
    position: relative;
    z-index: 2;
    transition: all .3s ease;
}
.pnode-circle.waiting  { background: var(--bg-hover);  border: 2px solid var(--border); }
.pnode-circle.active   { background: linear-gradient(135deg, var(--gold-dim), var(--gold)); border: none; animation: glow-pulse 1.8s ease-in-out infinite; }
.pnode-circle.done     { background: rgba(63,185,80,.15); border: 2px solid #2a6a3a; }

/* Pulse ring on active node */
.pnode-pulse {
    position: absolute;
    inset: 0;
    border-radius: 50%;
    border: 2px solid var(--gold);
    animation: pulse-ring 1.8s ease-out infinite;
    z-index: 1;
}

/* Connector line between nodes */
.pconnector {
    flex: 1;
    height: 3px;
    background: var(--border);
    border-radius: 2px;
    position: relative;
    overflow: hidden;
    margin: 0 4px;
    margin-bottom: 28px;  /* align with circle centres */
    z-index: 0;
}
/* Active flow particle */
.pconnector.flowing::after {
    content: '';
    position: absolute;
    top: 0; left: 0;
    width: 60%; height: 100%;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
    animation: flow-right 1.6s ease-in-out infinite;
}
/* Completed connector */
.pconnector.done-line {
    background: linear-gradient(90deg, #2a6a3a, var(--success));
}

.pnode-label {
    font-size: .75rem;
    font-weight: 600;
    color: var(--text-secondary);
    text-align: center;
    letter-spacing: .03em;
}
.pnode-label.active { color: var(--gold); }
.pnode-label.done   { color: var(--success); }

.pnode-status {
    font-size: .68rem;
    color: var(--text-muted);
    text-align: center;
    margin-top: -4px;
}
.pnode-status.active { color: var(--gold-light); }
.pnode-status.done   { color: var(--success); }

/* ── Bouncing dots loader ── */
.dots-loader {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 7px;
    padding: 8px 0 4px;
}
.dots-loader span {
    width: 9px; height: 9px;
    border-radius: 50%;
    background: var(--gold);
    animation: bounce-dots 1.2s ease-in-out infinite;
    display: inline-block;
}
.dots-loader span:nth-child(2) { animation-delay: .15s; }
.dots-loader span:nth-child(3) { animation-delay: .30s; }

/* ── Spinning ring loader ── */
.spin-loader {
    width: 44px; height: 44px;
    border-radius: 50%;
    border: 3px solid var(--border);
    border-top-color: var(--gold);
    animation: spin-ring .9s linear infinite;
    margin: 0 auto 14px;
}

/* ── Processing card (replaces static research-card during active state) ── */
.processing-card {
    background: var(--bg-raised);
    border: 1px solid var(--gold-dim);
    border-radius: var(--radius-md);
    padding: 32px 28px;
    text-align: center;
    animation: fade-in-up .4s ease both;
    box-shadow: 0 0 24px rgba(212,168,67,.08);
}
.processing-card .pc-title {
    font-family: 'Playfair Display', serif;
    font-size: 1.15rem;
    font-weight: 700;
    color: var(--text-primary);
    margin: 14px 0 6px;
}
.processing-card .pc-sub {
    font-size: .85rem;
    color: var(--text-secondary);
    margin-bottom: 16px;
}

/* ── Shimmer progress bar ── */
.progress-track {
    width: 100%;
    height: 4px;
    background: var(--border);
    border-radius: 2px;
    overflow: hidden;
    margin-top: 16px;
}
.progress-fill {
    height: 100%;
    border-radius: 2px;
    background: linear-gradient(90deg,
        var(--gold-dim) 0%,
        var(--gold) 40%,
        var(--gold-light) 60%,
        var(--gold) 100%);
    background-size: 300% 100%;
    animation: shimmer 1.8s linear infinite;
    width: 60%;
}

/* ── Step header active pulse ── */
.step-num-active { animation: glow-pulse 2s ease-in-out infinite; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────────────────────────────────────

def init_session():
    defaults = {
        "orchestrator": None,
        "messages": [],
        "pipeline_stage": PipelineStage.INTAKE,
        "report": None,
        "api_key": "",
        "pipeline_running": False,
        "active_tab": 0,
        "saved_app_id": None,
        "view_app_id": None,
        "page": "app",
        "agent_typing": False,
        # Phase 2 tracker
        "tracker_items": [],
        "tracker_refresh": 0,
        "appt_date": "",
        "appt_time": "",
        "appt_location": "",
        "appt_ref": "",
        "appt_portal_name": "",
        "appt_portal_url": "",
        "appt_saved": False,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _stage_to_tab(stage: PipelineStage) -> int:
    return {
        PipelineStage.INTAKE:      0,
        PipelineStage.RESEARCHING: 1,
        PipelineStage.ANALYSING:   2,
        PipelineStage.COMPLETE:    3,
        PipelineStage.ERROR:       3,
    }.get(stage, 0)


def _step_header(num: str, title: str, subtitle: str, state: str):
    """state: 'active' | 'done' | 'waiting'"""
    circle_cls = f"step-num-{state}"
    pill_cls   = f"pill-{state}"
    pill_txt   = {"active": "⏳ In Progress", "done": "✅ Complete", "waiting": "⏸ Waiting"}[state]
    icon       = {"active": num, "done": "✓", "waiting": num}[state]
    st.markdown(f"""
    <div class="step-header">
        <div class="step-num-circle {circle_cls}">{icon}</div>
        <div>
            <div class="step-title">{title}</div>
            <div class="step-subtitle">{subtitle}</div>
        </div>
        <span class="status-pill {pill_cls}">{pill_txt}</span>
    </div>""", unsafe_allow_html=True)


def render_chat(messages: list, show_typing: bool = False):
    html = '<div class="chat-wrap">'
    for msg in messages:
        content = msg["content"].replace("\n", "<br>").replace("**", "")
        if msg["role"] == "user":
            html += f'<div class="msg-row-user"><div class="msg-bubble bubble-user">{content}</div><div class="msg-avatar avatar-user">👤</div></div>'
        else:
            html += f'<div class="msg-row-agent"><div class="msg-avatar avatar-agent">🤖</div><div class="msg-bubble bubble-agent">{content}</div></div>'
    if show_typing:
        html += ('<div class="msg-row-agent">'
                 '<div class="msg-avatar avatar-agent">🤖</div>'
                 '<div class="msg-bubble bubble-agent typing-bubble">'
                 '<span class="typing-dot"></span>'
                 '<span class="typing-dot"></span>'
                 '<span class="typing-dot"></span>'
                 '</div></div>')
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


def render_pipeline_bar(stage: PipelineStage):
    """
    Animated pipeline connector bar — uses st.columns to avoid Streamlit
    markdown HTML-escaping issues with inline styles.
    """
    order = [PipelineStage.INTAKE, PipelineStage.RESEARCHING, PipelineStage.ANALYSING, PipelineStage.COMPLETE]
    cur_idx = order.index(stage) if stage in order else 0

    nodes = [
        {"icon": "💬", "label": "Intake",    "sub": "Collecting info"},
        {"icon": "🔍", "label": "Research",  "sub": "Fetching requirements"},
        {"icon": "📊", "label": "Analysis",  "sub": "Building report"},
    ]

    is_running = stage in [PipelineStage.RESEARCHING, PipelineStage.ANALYSING]
    bar_extra = ' running' if is_running else ''

    # Build each node HTML independently (no inline animation-delay to avoid escaping)
    def _node_html(i, node):
        if i < cur_idx:
            state = "done"
            icon  = "✓"
        elif i == cur_idx and stage != PipelineStage.COMPLETE:
            state = "active"
            icon  = node["icon"]
        else:
            state = "waiting"
            icon  = node["icon"]

        pulse = '<div class="pnode-pulse"></div>' if state == "active" else ""
        return (
            '<div class="pnode">'
            '<div class="pnode-icon-wrap">'
            + pulse +
            '<div class="pnode-circle ' + state + '">' + icon + '</div>'
            '</div>'
            '<div class="pnode-label ' + state + '">' + node["label"] + '</div>'
            '<div class="pnode-status ' + state + '">' + node["sub"] + '</div>'
            '</div>'
        )

    def _conn_html(i):
        if i == cur_idx and stage != PipelineStage.COMPLETE:
            cls = "pconnector flowing"
        elif i < cur_idx:
            cls = "pconnector done-line"
        else:
            cls = "pconnector"
        return '<div class="' + cls + '"></div>'

    parts = []
    for i, node in enumerate(nodes):
        parts.append(_node_html(i, node))
        if i < len(nodes) - 1:
            parts.append(_conn_html(i))

    html = '<div class="pipeline-bar' + bar_extra + '">' + "".join(parts) + '</div>'
    st.markdown(html, unsafe_allow_html=True)


def render_agent_logs(events, agent_label: str, panel_key: str):
    """
    Render a clean summary of agent activity — notes, flags, and sources only.
    No raw HTML in detail blocks; each section is rendered with st.markdown.
    """
    if not events:
        return

    st.markdown(
        f'<div class="section-label" style="margin-top:24px;">📝 {agent_label} — Summary</div>',
        unsafe_allow_html=True,
    )

    # Collect the meaningful content sections from events
    notes_sections: list[tuple[str, list[str]]] = []   # (section_title, [bullet, ...])
    sources: list[dict] = []
    broken_sources: list[dict] = []
    summary_lines: list[str] = []

    for ev in events:
        msg = ev.message
        detail = ev.detail or ""

        # Notes sections (nationality notes, residence notes, flags, tips)
        if ev.icon == "📝" and detail:
            items = [line.lstrip("• ").strip() for line in detail.split("\n") if line.strip()]
            notes_sections.append((msg, items))

        # Verified sources block (✅ icon from verify_sources)
        elif ev.icon == "✅" and "verified" in msg.lower() and detail:
            for line in detail.split("\n"):
                line = line.lstrip("• ").strip()
                if " — " in line:
                    parts = line.split(" — ", 1)
                    sources.append({"title": parts[0].strip(), "url": parts[1].strip()})

        # Legacy sources block (🔗 icon)
        elif ev.icon == "🔗" and detail:
            for line in detail.split("\n"):
                line = line.lstrip("• ").strip()
                if " — " in line:
                    parts = line.split(" — ", 1)
                    sources.append({"title": parts[0].strip(), "url": parts[1].strip()})

        # Broken sources warning from verify_sources
        elif ev.level == "warning" and "could not be reached" in msg and detail:
            for line in detail.split("\n"):
                line = line.lstrip("• ").strip()
                # Format: "Title (reason) — url"
                if " — " in line:
                    left, url = line.rsplit(" — ", 1)
                    # Strip reason in parens from title
                    title = left.split("(")[0].strip()
                    reason = left[left.find("(")+1:left.rfind(")")].strip() if "(" in left else ""
                    broken_sources.append({"title": title, "url": url, "reason": reason})

        # Embassy guidance
        elif ev.icon == "🏛️" and detail:
            summary_lines.append(("Embassy Guidance", detail.strip()))

        # Priority flags (warning level with detail, ⚠️ icon)
        elif ev.level == "warning" and detail and ev.icon == "⚠️" and "could not be reached" not in msg:
            items = [line.lstrip("• ").strip() for line in detail.split("\n") if line.strip()]
            notes_sections.append((msg, items))

        # ETA note
        elif ev.icon == "⚡" and detail:
            summary_lines.append(("ETA Note", detail.strip()))

    # ── Render summary lines (embassy guidance, ETA) ─────────────────────────────
    for label, text in summary_lines:
        st.markdown(
            f'<div class="info-banner banner-blue" style="margin:8px 0;">'
            f'<strong>{label}:</strong> {text}</div>',
            unsafe_allow_html=True,
        )

    # ── Render notes sections ───────────────────────────────────────────────
    for section_title, items in notes_sections:
        if not items:
            continue
        # Choose banner style based on content
        if "flag" in section_title.lower() or "warning" in section_title.lower():
            banner_cls = "banner-amber"
            prefix = "⚠️"
        elif "tip" in section_title.lower() or "practical" in section_title.lower():
            banner_cls = "banner-blue"
            prefix = "💡"
        else:
            banner_cls = "banner-blue"
            prefix = "📌"

        st.markdown(
            f'<div style="margin:10px 0 4px;">'
            f'<span style="font-size:.78rem;font-weight:700;letter-spacing:.07em;text-transform:uppercase;'
            f'color:var(--text-muted);">{prefix} {section_title}</span></div>',
            unsafe_allow_html=True,
        )
        for item in items:
            st.markdown(
                f'<div style="display:flex;gap:8px;align-items:flex-start;padding:5px 0;'
                f'border-bottom:1px solid var(--border-light);">'
                f'<span style="color:var(--gold);flex-shrink:0;margin-top:2px;">&#8250;</span>'
                f'<span style="font-size:.86rem;color:var(--text-primary);line-height:1.5;">{item}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )

    # ── Render sources ───────────────────────────────────────────────────────────
    if sources:
        st.markdown(
            '<div style="margin:14px 0 6px;">'
            '<span style="font-size:.78rem;font-weight:700;letter-spacing:.07em;text-transform:uppercase;'
            'color:var(--text-muted);">&#128279; Official Sources</span></div>',
            unsafe_allow_html=True,
        )
        links_html = ""
        for src in sources:
            if src.get("url"):
                links_html += f'<a class="source-link" href="{src["url"]}" target="_blank">↗ {src["title"]}</a>'
        st.markdown(links_html, unsafe_allow_html=True)

    # ── Render broken sources warning ────────────────────────────────────────
    if broken_sources:
        st.markdown(
            '<div style="margin:14px 0 6px;">'  
            '<span style="font-size:.78rem;font-weight:700;letter-spacing:.07em;text-transform:uppercase;'
            'color:#f59e0b;">⚠️ Inaccessible Links (removed from report)</span></div>',
            unsafe_allow_html=True,
        )
        for src in broken_sources:
            reason_text = f" — {src['reason']}" if src.get('reason') else ""
            st.markdown(
                f'<div style="display:flex;gap:8px;align-items:flex-start;padding:5px 0;'
                f'border-bottom:1px solid var(--border-light);">'
                f'<span style="color:#f59e0b;flex-shrink:0;margin-top:2px;">❌</span>'
                f'<span style="font-size:.83rem;color:var(--text-secondary);line-height:1.5;">'
                f'<strong style="color:var(--text-primary);">{src["title"]}</strong>'
                f'<span style="color:#f59e0b;">{reason_text}</span>'
                f'</span></div>',
                unsafe_allow_html=True,
            )


def render_report(report: VisaReport, key_prefix: str = "main"):
    dest = report.destination_country
    if report.schengen_main_country:
        dest = f"{report.destination_country} · {report.schengen_main_country}"

    st.markdown(f"""
    <div class="report-hero">
        <h2>🛂 {report.visa_type}</h2>
        <p>{report.applicant_nationality} national &nbsp;·&nbsp; Applying from {report.country_of_residence} &nbsp;·&nbsp; Destination: {dest}</p>
    </div>""", unsafe_allow_html=True)

    if report.eta_eligible and report.eta_note:
        st.markdown(f'<div class="info-banner banner-green">⚡ <strong>ETA Eligible:</strong> {report.eta_note}</div>', unsafe_allow_html=True)
    if report.embassy_guidance:
        st.markdown(f'<div class="info-banner banner-blue">🏛️ <strong>Embassy Selection:</strong> {report.embassy_guidance}</div>', unsafe_allow_html=True)
    if report.executive_summary:
        st.markdown(f'<div class="info-banner banner-blue">📋 <strong>Summary:</strong> {report.executive_summary}</div>', unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    fee_short  = report.fee.split("(")[0].strip()
    proc_short = report.processing_time.split(",")[0]
    with c1:
        st.markdown(f'<div class="metric-card"><div class="metric-label">💰 Visa Fee</div><div class="metric-value">{fee_short}</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="metric-card"><div class="metric-label">⏱ Processing Time</div><div class="metric-value">{proc_short}</div></div>', unsafe_allow_html=True)
    with c3:
        st.markdown(f'<div class="metric-card"><div class="metric-label">📅 Maximum Stay</div><div class="metric-value">{report.max_stay}</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown('<div class="section-label">📄 Required Documents</div>', unsafe_allow_html=True)
    for doc in report.mandatory_documents:
        with st.expander(f"✅  {doc.document}"):
            st.markdown(f'<div class="doc-detail">{doc.details}</div>', unsafe_allow_html=True)

    if report.optional_documents:
        st.markdown('<div class="section-label" style="margin-top:20px;">⭐ Recommended Documents</div>', unsafe_allow_html=True)
        for doc in report.optional_documents:
            with st.expander(f"⭐  {doc.document}"):
                st.markdown(f'<div class="doc-detail">{doc.details}</div>', unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">📋 Application Steps</div>', unsafe_allow_html=True)
    steps_html = ""
    for i, step in enumerate(report.application_steps, 1):
        text = step.split(":", 1)[-1].strip() if ":" in step else step
        steps_html += f'<div class="step-row"><div class="step-num">{i}</div><div class="step-text">{text}</div></div>'
    st.markdown(steps_html, unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    if report.supplementary_notes:
        st.markdown('<div class="section-label">🔍 Research Agent Findings</div>', unsafe_allow_html=True)
        for note in report.supplementary_notes:
            if note.startswith("**") and note.endswith(":**"):
                st.markdown(f"<p style='color:#d4a843;font-weight:700;margin:12px 0 4px;'>{note.strip('*').strip(':')}</p>", unsafe_allow_html=True)
            else:
                st.markdown(f"<p style='color:#8b949e;font-size:.87rem;margin:4px 0;'>{note}</p>", unsafe_allow_html=True)

    if report.key_notes:
        st.markdown('<div class="section-label" style="margin-top:20px;">⚠️ Important Notes</div>', unsafe_allow_html=True)
        for note in report.key_notes:
            st.markdown(f'<div class="info-banner banner-amber" style="margin:6px 0;">⚠️ {note}</div>', unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown('<div class="section-label">🔗 Official Sources</div>', unsafe_allow_html=True)
    links_html = ""
    for src in report.sources:
        if src.get("url") and src.get("title"):
            links_html += f'<a class="source-link" href="{src["url"]}" target="_blank">↗ {src["title"]}</a>'
    st.markdown(links_html + "<br><br>", unsafe_allow_html=True)

    # ── Phase 2: Application Forms ───────────────────────────────────────────
    if getattr(report, "forms", None):
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">📄 Application Forms</div>', unsafe_allow_html=True)
        mandatory_forms = [f for f in report.forms if f.mandatory]
        optional_forms  = [f for f in report.forms if not f.mandatory]
        for fm in mandatory_forms:
            col_info, col_btn = st.columns([0.72, 0.28])
            with col_info:
                st.markdown(
                    f'<div style="padding:10px 0;">'
                    f'<span style="font-size:.88rem;font-weight:600;color:var(--text-primary);">✅ {fm.title}</span><br>'
                    f'<span style="font-size:.8rem;color:var(--text-secondary);">{fm.description}</span>'
                    f'{f'<br><span style="font-size:.78rem;color:var(--text-muted);">💡 {fm.notes}</span>' if fm.notes else ""}'
                    f'</div>',
                    unsafe_allow_html=True,
                )
            with col_btn:
                st.markdown(
                    f'<a href="{fm.url}" target="_blank" class="source-link" '
                    f'style="display:inline-flex;margin-top:12px;padding:8px 14px;font-weight:600;">'
                    f'↗ {fm.format}</a>',
                    unsafe_allow_html=True,
                )
        if optional_forms:
            with st.expander("⭐ Optional / alternative forms"):
                for fm in optional_forms:
                    st.markdown(
                        f'<div style="padding:6px 0;border-bottom:1px solid var(--border-light);">'
                        f'<a class="source-link" href="{fm.url}" target="_blank">↗ {fm.title}</a>'
                        f'<span style="font-size:.8rem;color:var(--text-secondary);margin-left:8px;">{fm.description}</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

    # ── Phase 2: Document Upload Portal ───────────────────────────────────────
    if getattr(report, "upload_portal", None):
        p = report.upload_portal
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">☁️ Document Upload Portal</div>', unsafe_allow_html=True)
        st.markdown(
            f'<div style="background:var(--bg-raised);border:1px solid var(--border);'
            f'border-left:4px solid var(--accent);border-radius:var(--radius-md);padding:18px 20px;">'
            f'<div style="font-size:.9rem;font-weight:700;color:var(--text-primary);margin-bottom:6px;">{p.name}</div>'
            f'<div style="display:flex;flex-wrap:wrap;gap:8px;margin-bottom:10px;">'
            f'<span class="visa-pill">📁 Formats: {", ".join(p.accepted_formats)}</span>'
            f'<span class="visa-pill">📦 Max: {p.max_file_mb}MB per file</span>'
            f'<span class="visa-pill">🕐 {p.upload_timing}</span>'
            f'</div>'
            f'<div style="font-size:.84rem;color:var(--text-secondary);margin-bottom:12px;">{p.notes}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
        col_portal, col_login = st.columns(2)
        with col_portal:
            st.markdown(f'<a href="{p.url}" target="_blank" class="source-link" style="padding:8px 14px;font-weight:600;">↗ Open upload portal</a>', unsafe_allow_html=True)
        with col_login:
            if p.login_instructions:
                with st.expander("📋 How to upload — step by step"):
                    for i, step in enumerate(p.login_instructions, 1):
                        st.markdown(f'<div class="step-row"><div class="step-num">{i}</div><div class="step-text">{step}</div></div>', unsafe_allow_html=True)

    # ── Phase 2: Appointment Portals ──────────────────────────────────────────
    if getattr(report, "appointment_portals", None):
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">📅 Appointment Booking</div>', unsafe_allow_html=True)
        for portal in report.appointment_portals:
            st.markdown(
                f'<div style="background:var(--bg-raised);border:1px solid var(--border);'
                f'border-left:4px solid var(--gold);border-radius:var(--radius-md);padding:18px 20px;margin-bottom:10px;">'
                f'<div style="font-size:.9rem;font-weight:700;color:var(--text-primary);margin-bottom:6px;">{portal.name}</div>'
                f'<div style="display:flex;flex-wrap:wrap;gap:8px;margin-bottom:10px;">'
                f'<span class="visa-pill">⏳ ~{portal.avg_wait_weeks} week wait</span>'
                f'</div>'
                f'<div style="font-size:.84rem;color:var(--text-secondary);margin-bottom:12px;">{portal.notes}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
            col_book, col_locs = st.columns(2)
            with col_book:
                st.markdown(f'<a href="{portal.booking_url}" target="_blank" class="source-link" style="padding:8px 14px;font-weight:600;">↗ Book appointment</a>', unsafe_allow_html=True)
            with col_locs:
                if portal.locations_info_url:
                    st.markdown(f'<a href="{portal.locations_info_url}" target="_blank" class="source-link" style="padding:8px 14px;">↗ Find nearest centre</a>', unsafe_allow_html=True)
            if portal.booking_steps:
                with st.expander("📋 How to book — step by step"):
                    for i, step in enumerate(portal.booking_steps, 1):
                        st.markdown(f'<div class="step-row"><div class="step-num">{i}</div><div class="step-text">{step}</div></div>', unsafe_allow_html=True)

    # ── Phase 2: Status Tracking ──────────────────────────────────────────────
    if getattr(report, "tracking_url", None):
        st.markdown(
            f'<div class="info-banner banner-blue" style="margin-top:16px;">'
            f'🔎 <strong>Track your application:</strong> '
            f'<a href="{report.tracking_url}" target="_blank" style="color:var(--accent);">{report.tracking_url}</a>'
            f'{f" — {report.tracking_instructions}" if report.tracking_instructions else ""}'
            f'</div>',
            unsafe_allow_html=True,
        )

    _base = report.destination_country.replace(' ', '_').lower()
    _k    = f"{key_prefix}_{_base}"

    st.markdown('<div class="section-label" style="margin-top:8px;">⬇️ Export Report</div>', unsafe_allow_html=True)
    col_md, col_pdf, col_docx = st.columns(3)

    with col_md:
        st.download_button(
            label="📄 Markdown (.md)",
            data=_build_markdown(report),
            file_name=f"visa_report_{_base}.md",
            mime="text/markdown",
            use_container_width=True,
            key=f"dl_md_{_k}",
        )
    with col_pdf:
        try:
            pdf_bytes = _build_pdf(report)
            st.download_button(
                label="📕 PDF (.pdf)",
                data=pdf_bytes,
                file_name=f"visa_report_{_base}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key=f"dl_pdf_{_k}",
            )
        except Exception as e:
            st.button(f"📕 PDF (error: {e})", disabled=True, use_container_width=True, key=f"dl_pdf_err_{_k}")
    with col_docx:
        try:
            docx_bytes = _build_docx(report)
            st.download_button(
                label="📘 Word (.docx)",
                data=docx_bytes,
                file_name=f"visa_report_{_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"dl_docx_{_k}",
            )
        except Exception as e:
            st.button(f"📘 Word (error: {e})", disabled=True, use_container_width=True, key=f"dl_docx_err_{_k}")


def _build_markdown(report: VisaReport) -> str:
    lines = [
        f"# {report.visa_type} — Requirements Report", "",
        f"**Applicant:** {report.applicant_nationality} national, residing in {report.country_of_residence}",
        f"**Destination:** {report.destination_country}", f"**Purpose:** {report.purpose_of_visit}", "",
        "---", "", "## Summary", report.executive_summary, "",
        "## Visa Details",
        f"- **Fee:** {report.fee}",
        f"- **Processing Time:** {report.processing_time}",
        f"- **Maximum Stay:** {report.max_stay}",
        f"- **Apply Online:** {report.apply_url}", "",
        "## Required Documents",
    ]
    for d in report.mandatory_documents:
        lines.append(f"- **{d.document}:** {d.details}")
    lines += ["", "## Recommended Documents"]
    for d in report.optional_documents:
        lines.append(f"- **{d.document}:** {d.details}")
    lines += ["", "## Application Steps"]
    for i, s in enumerate(report.application_steps, 1):
        lines.append(f"{i}. {s}")
    if report.supplementary_notes:
        lines += ["", "## Research Findings"]
        lines.extend(report.supplementary_notes)
    lines += ["", "## Key Notes"]
    for n in report.key_notes:
        lines.append(f"- {n}")
    lines += ["", "## Official Sources"]
    for s in report.sources:
        if s.get("url"):
            lines.append(f"- [{s['title']}]({s['url']})")
    lines += ["", "---", "*Generated by Visa Application Multi-Agent System — Phase 1*"]
    return "\n".join(lines)


def _build_pdf(report: VisaReport) -> bytes:
    """Generate a styled PDF from the VisaReport using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=18*mm, bottomMargin=18*mm,
    )

    GOLD   = colors.HexColor("#d4a843")
    DARK   = colors.HexColor("#1c2333")
    MUTED  = colors.HexColor("#6e7681")
    WHITE  = colors.white
    GREEN  = colors.HexColor("#3fb950")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"],
        fontSize=20, textColor=GOLD, spaceAfter=4, spaceBefore=0, leading=26)
    sub_style   = ParagraphStyle("Sub", parent=styles["Normal"],
        fontSize=10, textColor=MUTED, spaceAfter=12)
    h2_style    = ParagraphStyle("H2", parent=styles["Heading2"],
        fontSize=13, textColor=GOLD, spaceBefore=14, spaceAfter=4, borderPad=0)
    body_style  = ParagraphStyle("Body", parent=styles["Normal"],
        fontSize=9.5, textColor=colors.HexColor("#2c3e50"), leading=14, spaceAfter=4)
    label_style = ParagraphStyle("Label", parent=styles["Normal"],
        fontSize=8.5, textColor=MUTED, spaceAfter=2, fontName="Helvetica-Bold")
    note_style  = ParagraphStyle("Note", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#555"), leading=13,
        leftIndent=10, spaceAfter=3)
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"],
        fontSize=8, textColor=MUTED, alignment=TA_CENTER)

    dest = report.destination_country
    if report.schengen_main_country:
        dest += f" · {report.schengen_main_country}"

    story = []
    story.append(Paragraph(f"{report.visa_type} — Requirements Report", title_style))
    story.append(Paragraph(
        f"{report.applicant_nationality} national · Applying from {report.country_of_residence} · Destination: {dest}",
        sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceAfter=8))

    if report.executive_summary:
        story.append(Paragraph("Summary", h2_style))
        story.append(Paragraph(report.executive_summary, body_style))

    story.append(Paragraph("Visa Details", h2_style))
    for label, val in [("Fee", report.fee), ("Processing Time", report.processing_time),
                       ("Maximum Stay", report.max_stay)]:
        story.append(Paragraph(f"<b>{label}:</b> {val}", body_style))
    if report.apply_url:
        story.append(Paragraph(f"<b>Apply Online:</b> {report.apply_url}", body_style))

    story.append(Paragraph("Required Documents", h2_style))
    for d in report.mandatory_documents:
        story.append(Paragraph(f"✓ <b>{d.document}</b>", body_style))
        story.append(Paragraph(d.details, note_style))

    if report.optional_documents:
        story.append(Paragraph("Recommended Documents", h2_style))
        for d in report.optional_documents:
            story.append(Paragraph(f"★ <b>{d.document}</b>", body_style))
            story.append(Paragraph(d.details, note_style))

    story.append(Paragraph("Application Steps", h2_style))
    for i, step in enumerate(report.application_steps, 1):
        text = step.split(":", 1)[-1].strip() if ":" in step else step
        story.append(Paragraph(f"{i}. {text}", body_style))

    if report.key_notes:
        story.append(Paragraph("Important Notes", h2_style))
        for note in report.key_notes:
            story.append(Paragraph(f"⚠ {note}", note_style))

    if report.sources:
        story.append(Paragraph("Official Sources", h2_style))
        for src in report.sources:
            if src.get("url") and src.get("title"):
                story.append(Paragraph(f"→ <a href='{src['url']}'>{src['title']}</a>", body_style))

    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=0.5, color=MUTED))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Generated by Visa Application Assistant — Phase 2", footer_style))

    doc.build(story)
    return buf.getvalue()


def _build_docx(report: VisaReport) -> bytes:
    """Generate a styled DOCX from the VisaReport using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GOLD  = RGBColor(0xD4, 0xA8, 0x43)
    DARK  = RGBColor(0x1C, 0x23, 0x33)
    MUTED = RGBColor(0x6E, 0x76, 0x81)

    document = Document()

    # Page margins
    for section in document.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def _heading(text, level=1, color=GOLD):
        p = document.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
        return p

    def _body(text, bold=False, color=None):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
        return p

    def _bullet(text, bold_prefix=None):
        p = document.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix + ": ")
            r.font.bold = True
            r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        return p

    dest = report.destination_country
    if report.schengen_main_country:
        dest += f" · {report.schengen_main_country}"

    _heading(f"{report.visa_type} — Requirements Report", level=1)
    _body(f"{report.applicant_nationality} national · Applying from {report.country_of_residence} · Destination: {dest}",
          color=MUTED)
    document.add_paragraph()

    if report.executive_summary:
        _heading("Summary", level=2)
        _body(report.executive_summary)

    _heading("Visa Details", level=2)
    for label, val in [("Fee", report.fee), ("Processing Time", report.processing_time),
                       ("Maximum Stay", report.max_stay)]:
        p = document.add_paragraph()
        r = p.add_run(label + ": "); r.font.bold = True; r.font.size = Pt(10)
        r2 = p.add_run(val); r2.font.size = Pt(10)
    if report.apply_url:
        p = document.add_paragraph()
        r = p.add_run("Apply Online: "); r.font.bold = True; r.font.size = Pt(10)
        r2 = p.add_run(report.apply_url); r2.font.size = Pt(10)

    _heading("Required Documents", level=2)
    for d in report.mandatory_documents:
        _bullet(d.details, bold_prefix=d.document)

    if report.optional_documents:
        _heading("Recommended Documents", level=2)
        for d in report.optional_documents:
            _bullet(d.details, bold_prefix=d.document)

    _heading("Application Steps", level=2)
    for i, step in enumerate(report.application_steps, 1):
        text = step.split(":", 1)[-1].strip() if ":" in step else step
        p = document.add_paragraph(style="List Number")
        r = p.add_run(text); r.font.size = Pt(10)

    if report.key_notes:
        _heading("Important Notes", level=2)
        for note in report.key_notes:
            _bullet(note)

    if report.sources:
        _heading("Official Sources", level=2)
        for src in report.sources:
            if src.get("url") and src.get("title"):
                _bullet(f"{src['title']} — {src['url']}")

    document.add_paragraph()
    p = document.add_paragraph("Generated by Visa Application Assistant — Phase 2")
    p.runs[0].font.color.rgb = MUTED
    p.runs[0].font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    init_session()

    with st.sidebar:
        st.markdown("""
        <div class="sidebar-logo">
            <div class="logo-icon">🛂</div>
            <div class="logo-title">Visa Assistant</div>
            <div class="logo-sub">AI-Powered · Phase 2</div>
        </div>""", unsafe_allow_html=True)

        api_key = st.text_input(
            "OpenAI API Key", type="password",
            value=st.session_state.api_key, placeholder="sk-…",
            help="Your key is never stored permanently.",
        )
        if api_key:
            st.session_state.api_key = api_key

        model = st.selectbox("Model", ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo"], index=0)

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Supported Visas</div>', unsafe_allow_html=True)
        st.markdown('<div class="visa-pill">🇬🇧 UK Standard Visitor Visa</div>', unsafe_allow_html=True)
        st.markdown('<div class="visa-pill">🇪🇺 Schengen Short-Stay (Type C)</div>', unsafe_allow_html=True)
        st.markdown("<hr>", unsafe_allow_html=True)

        if st.button("🔄 New Application", use_container_width=True):
            if st.session_state.orchestrator:
                st.session_state.orchestrator.reset()
            for k in ["orchestrator", "messages", "report", "pipeline_running", "saved_app_id"]:
                st.session_state[k] = None if k in ("orchestrator", "saved_app_id") else ([] if k == "messages" else (None if k == "report" else False))
            for k in ["tracker_items", "appt_date", "appt_time", "appt_location", "appt_ref", "appt_saved"]:
                st.session_state[k] = [] if k == "tracker_items" else ""
            st.session_state.appt_saved = False
            st.session_state.pipeline_stage = PipelineStage.INTAKE
            st.session_state.active_tab = 0
            st.rerun()

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="disclaimer">⚠️ General information only. Always verify with the official embassy or consulate before applying.</div>', unsafe_allow_html=True)

    # ── Hero ──────────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="page-hero">
        <h1>🛂 Visa Application Assistant</h1>
        <p>Multi-Agent AI System &nbsp;·&nbsp; <span class="gold-line">Phase 2</span> — Research · Forms · Portals · Progress Tracker</p>
    </div>""", unsafe_allow_html=True)

    # ── Animated pipeline bar ─────────────────────────────────────────────────
    render_pipeline_bar(st.session_state.pipeline_stage)

    # ── API key guard ─────────────────────────────────────────────────────────
    if not st.session_state.api_key:
        st.markdown('<div class="info-banner banner-amber">👈 <strong>Enter your OpenAI API key</strong> in the sidebar to get started.</div>', unsafe_allow_html=True)
        st.stop()

    # ── Orchestrator init ─────────────────────────────────────────────────────
    if st.session_state.orchestrator is None:
        st.session_state.orchestrator = VisaOrchestrator(
            openai_api_key=st.session_state.api_key,
            model=model,
        )
        st.session_state.messages.append({
            "role": "assistant",
            "content": (
                "Hello! 👋 I'm your Visa Application Assistant. "
                "I'll help you understand the requirements for your UK or Schengen visa application.\n\n"
                "Let's start — what is your nationality (the country of your passport)?"
            ),
        })

    stage = st.session_state.pipeline_stage

    # ── Tab state mapping ─────────────────────────────────────────────────────
    # Determine state for each tab header
    def tab_state(tab_stage: PipelineStage) -> str:
        order = [PipelineStage.INTAKE, PipelineStage.RESEARCHING, PipelineStage.ANALYSING, PipelineStage.COMPLETE]
        cur_idx = order.index(stage) if stage in order else 0
        tab_idx = order.index(tab_stage)
        if cur_idx == tab_idx:
            return "active"
        elif cur_idx > tab_idx:
            return "done"
        return "waiting"

    # ── Six tabs ──────────────────────────────────────────────────────────────
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "💬  Step 1 · Intake",
        "🔍  Step 2 · Research",
        "📊  Step 3 · Analysis",
        "📋  Step 4 · Report",
        "✅  Step 5 · Track Progress",
        "🗂️  History",
    ])

    # ── TAB 1: INTAKE ─────────────────────────────────────────────────────────
    with tab1:
        _step_header("1", "Intake Agent", "Collecting your travel information", tab_state(PipelineStage.INTAKE))

        if stage == PipelineStage.INTAKE:
            if st.session_state.agent_typing:
                # Show typing bubble while agent is generating reply
                render_chat(st.session_state.messages, show_typing=True)
                orch = st.session_state.orchestrator
                last_user = next((m["content"] for m in reversed(st.session_state.messages) if m["role"] == "user"), "")
                reply, new_stage = orch.chat(last_user)
                st.session_state.pipeline_stage = new_stage
                st.session_state.agent_typing = False
                display = reply.split("```json")[0].strip() if "```json" in reply else reply
                if display:
                    st.session_state.messages.append({"role": "assistant", "content": display})
                st.rerun()
            else:
                render_chat(st.session_state.messages, show_typing=False)
                with st.form("chat_form", clear_on_submit=True):
                    col_in, col_btn = st.columns([0.82, 0.18])
                    with col_in:
                        user_input = st.text_input("msg", placeholder="Type your answer here…", label_visibility="collapsed")
                    with col_btn:
                        submitted = st.form_submit_button("Send →", use_container_width=True)

                if submitted and user_input.strip():
                    st.session_state.messages.append({"role": "user", "content": user_input})
                    st.session_state.agent_typing = True
                    st.rerun()

        elif stage in [PipelineStage.RESEARCHING, PipelineStage.ANALYSING, PipelineStage.COMPLETE]:
            # Show read-only chat history
            render_chat(st.session_state.messages)
            st.markdown('<div class="info-banner banner-green" style="margin-top:16px;">✅ Intake complete — your profile has been collected. View the Research tab to continue.</div>', unsafe_allow_html=True)

    # ── TAB 2: RESEARCH ───────────────────────────────────────────────────────
    with tab2:
        _step_header("2", "Research Agent", "Retrieving official visa requirements", tab_state(PipelineStage.RESEARCHING))

        if stage == PipelineStage.INTAKE:
            st.markdown("""
            <div class="research-card">
                <div class="rc-icon">🔒</div>
                <div class="rc-title">Waiting for Intake</div>
                <div class="rc-sub">Complete the intake conversation in Step 1 first.</div>
            </div>""", unsafe_allow_html=True)

        elif stage == PipelineStage.RESEARCHING and not st.session_state.pipeline_running:
            st.session_state.pipeline_running = True
            st.markdown("""
            <div class="processing-card">
                <div class="spin-loader"></div>
                <div class="pc-title">Research Agent Starting…</div>
                <div class="pc-sub">Loading the visa requirements knowledge base.</div>
                <div class="dots-loader"><span></span><span></span><span></span></div>
                <div class="progress-track"><div class="progress-fill"></div></div>
            </div>""", unsafe_allow_html=True)
            st.rerun()

        elif stage == PipelineStage.RESEARCHING and st.session_state.pipeline_running:
            st.markdown("""
            <div class="processing-card">
                <div class="spin-loader"></div>
                <div class="pc-title">Research Agent Working…</div>
                <div class="pc-sub">Retrieving official visa requirements and verifying sources.</div>
                <div class="dots-loader"><span></span><span></span><span></span></div>
                <div class="progress-track"><div class="progress-fill"></div></div>
            </div>""", unsafe_allow_html=True)
            with st.spinner(""):
                st.session_state.pipeline_stage = PipelineStage.ANALYSING
                st.rerun()

        elif stage in [PipelineStage.ANALYSING, PipelineStage.COMPLETE]:
            profile = st.session_state.orchestrator.profile if st.session_state.orchestrator else None
            st.markdown("""
            <div class="research-card" style="text-align:left;">
                <div style="font-size:1.5rem;margin-bottom:10px;">✅</div>
                <div class="rc-title" style="text-align:left;margin-bottom:12px;">Research Complete</div>""", unsafe_allow_html=True)
            if profile:
                st.markdown(f"""
                <div style="display:flex;flex-wrap:wrap;gap:10px;margin-top:8px;">
                    <span class="visa-pill">🌍 Nationality: {profile.nationality}</span>
                    <span class="visa-pill">🏠 Residence: {profile.country_of_residence}</span>
                    <span class="visa-pill">✈️ Destination: {profile.destination_country}</span>
                    <span class="visa-pill">🎯 Purpose: {profile.purpose_of_visit}</span>
                </div>""", unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
            st.markdown('<div class="info-banner banner-blue" style="margin-top:16px;">📊 Requirements retrieved. The Analysis Agent is now generating your personalised report — view Step 3.</div>', unsafe_allow_html=True)
            # ── Research Agent Activity Log ──────────────────────────────────
            orch = st.session_state.orchestrator
            if orch:
                research_logs = orch.logs_for("research")
                render_agent_logs(research_logs, "Research Agent", "research_log")

    # ── TAB 3: ANALYSIS ───────────────────────────────────────────────────────
    with tab3:
        _step_header("3", "Analysis Agent", "Generating your personalised requirements report", tab_state(PipelineStage.ANALYSING))

        if stage in [PipelineStage.INTAKE, PipelineStage.RESEARCHING]:
            st.markdown("""
            <div class="research-card">
                <div class="rc-icon">🔒</div>
                <div class="rc-title">Waiting for Research</div>
                <div class="rc-sub">The Research Agent must complete Step 2 first.</div>
            </div>""", unsafe_allow_html=True)
            if stage == PipelineStage.RESEARCHING:
                st.markdown("""
                <div class="processing-card" style="margin-top:16px;">
                    <div class="spin-loader"></div>
                    <div class="pc-title">Research Agent is running…</div>
                    <div class="pc-sub">Switch to Step 2 to monitor progress.</div>
                    <div class="dots-loader"><span></span><span></span><span></span></div>
                </div>""", unsafe_allow_html=True)

        elif stage == PipelineStage.ANALYSING:
            st.markdown("""
            <div class="processing-card">
                <div class="spin-loader"></div>
                <div class="pc-title">Analysis Agent Working…</div>
                <div class="pc-sub">Synthesising requirements into your personalised report. This may take 15–30 seconds.</div>
                <div class="dots-loader"><span></span><span></span><span></span></div>
                <div class="progress-track"><div class="progress-fill"></div></div>
            </div>""", unsafe_allow_html=True)
            with st.spinner(""):
                try:
                    report = st.session_state.orchestrator.run_research_and_analysis()
                    st.session_state.report = report
                    st.session_state.pipeline_stage = PipelineStage.COMPLETE
                    st.session_state.pipeline_running = False
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": f"✅ Your {report.visa_type} requirements report is ready. View Step 4 for the full report.",
                    })
                    # ── Auto-save to database ──────────────────────────────
                    try:
                        app_id = db.save_application(
                            profile=st.session_state.orchestrator.profile,
                            report=report,
                            chat_history=st.session_state.messages,
                        )
                        st.session_state.saved_app_id = app_id
                    except Exception as save_err:
                        st.session_state.saved_app_id = None
                        print(f"[DB] Save failed: {save_err}")
                except Exception as e:
                    st.session_state.pipeline_stage = PipelineStage.ERROR
                    st.markdown(f'<div class="info-banner banner-red">❌ <strong>Error:</strong> {e}</div>', unsafe_allow_html=True)
            st.rerun()

        elif stage == PipelineStage.COMPLETE:
            st.markdown("""
            <div class="research-card" style="text-align:left;">
                <div style="font-size:1.5rem;margin-bottom:10px;">✅</div>
                <div class="rc-title" style="text-align:left;">Analysis Complete</div>
            </div>""", unsafe_allow_html=True)
            st.markdown('<div class="info-banner banner-green" style="margin-top:16px;">📋 Your report is ready — view Step 4 for the full requirements, document checklist, and application steps.</div>', unsafe_allow_html=True)
            # ── Analysis Agent Activity Log ────────────────────────────────
            orch = st.session_state.orchestrator
            if orch:
                analysis_logs = orch.logs_for("analysis")
                render_agent_logs(analysis_logs, "Analysis Agent", "analysis_log")

    # ── TAB 4: REPORT ─────────────────────────────────────────────────────────
    with tab4:
        _step_header("4", "Requirements Report", "Your personalised visa requirements and checklist", tab_state(PipelineStage.COMPLETE))

        if stage != PipelineStage.COMPLETE or not st.session_state.report:
            st.markdown("""
            <div class="research-card">
                <div class="rc-icon">🔒</div>
                <div class="rc-title">Report Not Yet Generated</div>
                <div class="rc-sub">Complete Steps 1–3 to generate your personalised report.</div>
            </div>""", unsafe_allow_html=True)

            # Progress indicator
            steps_done = _stage_to_tab(stage)
            st.markdown(f"""
            <div style="margin-top:20px;">
                <div class="section-label">Progress</div>
                <div style="display:flex;gap:8px;align-items:center;">
                    {''.join([
                        f'<span style="padding:5px 14px;border-radius:20px;font-size:.8rem;font-weight:600;background:{"rgba(63,185,80,.12)" if i < steps_done else ("rgba(212,168,67,.12)" if i == steps_done else "var(--bg-hover)")};color:{"#3fb950" if i < steps_done else ("#d4a843" if i == steps_done else "#6e7681")};border:1px solid {"#1a5c2a" if i < steps_done else ("#8a6a20" if i == steps_done else "var(--border)")}">{["Intake","Research","Analysis","Report"][i]}</span>'
                        for i in range(4)
                    ])}
                </div>
            </div>""", unsafe_allow_html=True)
        else:
            # Show save confirmation badge
            if st.session_state.saved_app_id:
                st.markdown(
                    f'<div class="info-banner banner-green" style="margin-bottom:16px;">'
                    f'💾 <strong>Saved to history</strong> — Application #{st.session_state.saved_app_id} stored in database. '
                    f'View it anytime in the <strong>History</strong> tab.</div>',
                    unsafe_allow_html=True,
                )
            render_report(st.session_state.report, key_prefix="report_tab")

    # ── TAB 5: TRACKER ────────────────────────────────────────────────────────
    with tab5:
        _step_header("5", "Track Progress", "Your step-by-step application checklist", tab_state(PipelineStage.COMPLETE))

        if stage != PipelineStage.COMPLETE or not st.session_state.saved_app_id:
            st.markdown("""
            <div class="research-card">
                <div class="rc-icon">🔒</div>
                <div class="rc-title">Tracker not yet available</div>
                <div class="rc-sub">Complete Steps 1–4 to generate your personalised tracker.</div>
            </div>""", unsafe_allow_html=True)
        else:
            app_id = st.session_state.saved_app_id
            progress = db.get_progress(app_id)
            tracker_items = db.get_tracker_items(app_id)
            report = st.session_state.report

            # ── Progress bar ──────────────────────────────────────────────
            pct = progress["pct_complete"]
            bar_color = "#3fb950" if pct == 100 else "#d4a843"
            st.markdown(
                f'<div style="margin-bottom:20px;">'
                f'<div style="display:flex;justify-content:space-between;margin-bottom:6px;">'
                f'<span style="font-size:.82rem;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:var(--text-muted);">Overall Progress</span>'
                f'<span style="font-size:.88rem;font-weight:700;color:{bar_color};">{pct}% complete</span>'
                f'</div>'
                f'<div style="background:var(--bg-hover);border-radius:4px;height:8px;overflow:hidden;">'
                f'<div style="background:{bar_color};width:{pct}%;height:100%;border-radius:4px;transition:width .4s;"></div>'
                f'</div>'
                f'<div style="display:flex;gap:16px;margin-top:8px;font-size:.78rem;color:var(--text-muted);">'
                f'<span>✅ Done: {progress["done"]}</span>'
                f'<span>🔄 In progress: {progress["in_progress"]}</span>'
                f'<span>⏳ Pending: {progress["pending"]}</span>'
                f'{f'<span style="color:#f59e0b;">⚠️ Blocked: {progress["blocked"]}</span>' if progress["blocked"] else ""}'
                f'</div></div>',
                unsafe_allow_html=True,
            )

            # ── Appointment section ───────────────────────────────────────
            appointment = db.get_appointment(app_id)
            has_appt = appointment and appointment.get("reference_number")
            portal = report.appointment_portals[0] if getattr(report, "appointment_portals", None) else None

            st.markdown('<div class="section-label">📅 Appointment</div>', unsafe_allow_html=True)
            if has_appt:
                st.markdown(
                    f'<div class="info-banner banner-green">✅ <strong>Appointment booked</strong> — '
                    f'{appointment["portal_name"]} · {appointment["appointment_date"]} {appointment["appointment_time"]} · '
                    f'Ref: <strong>{appointment["reference_number"]}</strong></div>',
                    unsafe_allow_html=True,
                )
                # ICS download
                if appointment.get("ics_content"):
                    dest_slug = report.destination_country.replace(" ", "_").lower()
                    st.download_button(
                        "📅 Download calendar event (.ics)",
                        data=appointment["ics_content"],
                        file_name=f"visa_appointment_{dest_slug}.ics",
                        mime="text/calendar",
                        use_container_width=False,
                        key=f"ics_dl_{app_id}",
                    )
            else:
                # Booking guidance card
                if portal:
                    st.markdown(
                        f'<div style="background:var(--bg-raised);border:1px solid var(--border);'
                        f'border-left:4px solid var(--gold);border-radius:var(--radius-md);padding:16px 20px;margin-bottom:12px;">'
                        f'<div style="font-size:.88rem;font-weight:700;color:var(--text-primary);margin-bottom:4px;">{portal.name}</div>'
                        f'<span class="visa-pill" style="margin-bottom:8px;display:inline-block;">⏳ ~{portal.avg_wait_weeks} week wait</span>'
                        f'<div style="margin-top:8px;">'
                        f'<a href="{portal.booking_url}" target="_blank" class="source-link" style="padding:7px 14px;font-weight:600;">↗ Book appointment now</a>'
                        f'</div></div>',
                        unsafe_allow_html=True,
                    )

                # Reference number entry form
                st.markdown('<div style="margin-top:4px;margin-bottom:4px;font-size:.82rem;font-weight:700;color:var(--text-muted);letter-spacing:.06em;text-transform:uppercase;">Record your booking</div>', unsafe_allow_html=True)
                with st.form(f"appt_form_{app_id}", clear_on_submit=False):
                    portal_name_val = portal.name if portal else ""
                    portal_url_val  = portal.booking_url if portal else ""
                    col_a, col_b = st.columns(2)
                    with col_a:
                        appt_date = st.text_input("Appointment date (YYYY-MM-DD)", value=st.session_state.appt_date, placeholder="2025-06-15")
                        appt_loc  = st.text_input("Application centre / location", value=st.session_state.appt_location, placeholder="VFS Global Abu Dhabi, UAE")
                    with col_b:
                        appt_time = st.text_input("Time (HH:MM)", value=st.session_state.appt_time, placeholder="10:30")
                        appt_ref  = st.text_input("Reference number ★", value=st.session_state.appt_ref, placeholder="VFSAE123456789")
                    appt_notes = st.text_area("Notes (optional)", height=68, placeholder="e.g. bring 2 passport photos, fee paid")
                    save_appt = st.form_submit_button("💾 Save appointment & generate calendar", use_container_width=True)

                    if save_appt:
                        if not appt_ref.strip():
                            st.warning("Enter the reference number from your booking confirmation.")
                        else:
                            ics = generate_ics(
                                application_id=app_id,
                                portal_name=portal_name_val or appt_loc,
                                appointment_date=appt_date,
                                appointment_time=appt_time,
                                location=appt_loc,
                                reference_number=appt_ref,
                                visa_type=report.visa_type,
                            )
                            db.save_appointment(
                                app_id=app_id,
                                portal_name=portal_name_val or appt_loc,
                                portal_url=portal_url_val,
                                appointment_date=appt_date,
                                appointment_time=appt_time,
                                location=appt_loc,
                                reference_number=appt_ref,
                                confirmation_notes=appt_notes,
                                ics_content=ics,
                            )
                            st.session_state.appt_ref  = appt_ref
                            st.session_state.appt_date = appt_date
                            st.session_state.appt_saved = True
                            st.rerun()

            # ── Checklist ─────────────────────────────────────────────────
            st.markdown("<hr>", unsafe_allow_html=True)
            st.markdown('<div class="section-label">✅ Checklist</div>', unsafe_allow_html=True)

            STATUS_OPTIONS = [
                ItemStatus.PENDING,
                ItemStatus.IN_PROGRESS,
                ItemStatus.DONE,
                ItemStatus.BLOCKED,
                ItemStatus.NOT_NEEDED,
            ]
            STATUS_DISPLAY = {s: f"{STATUS_ICONS[s]} {STATUS_LABELS[s]}" for s in STATUS_OPTIONS}

            # Group by item type for cleaner UI
            type_order = [
                (ItemType.DOCUMENT,    "📄 Documents"),
                (ItemType.STEP,        "📋 Application steps"),
                (ItemType.APPOINTMENT, "📅 Appointment"),
                (ItemType.SUBMISSION,  "📤 Submission"),
                (ItemType.DECISION,    "🔎 Tracking & decision"),
            ]
            items_by_type = {}
            for item in tracker_items:
                items_by_type.setdefault(item["item_type"], []).append(item)

            for itype, group_label in type_order:
                group = items_by_type.get(itype, [])
                if not group:
                    continue
                st.markdown(f'<div style="font-size:.78rem;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:var(--text-muted);margin:18px 0 8px;">{group_label}</div>', unsafe_allow_html=True)
                for item in group:
                    with st.container():
                        icon = STATUS_ICONS.get(item["status"], "⏳")
                        col_icon, col_info, col_status, col_link = st.columns([0.05, 0.52, 0.25, 0.18])
                        with col_icon:
                            st.markdown(f'<div style="font-size:1.1rem;padding-top:6px;">{icon}</div>', unsafe_allow_html=True)
                        with col_info:
                            st.markdown(
                                f'<div style="padding:4px 0;">'
                                f'<div style="font-size:.88rem;font-weight:600;color:var(--text-primary);">{item["label"]}</div>'
                                f'{f'<div style="font-size:.78rem;color:var(--text-secondary);margin-top:2px;">{item["details"]}</div>' if item.get("details") else ""}'
                                f'</div>',
                                unsafe_allow_html=True,
                            )
                        with col_status:
                            new_status = st.selectbox(
                                "status",
                                options=STATUS_OPTIONS,
                                index=STATUS_OPTIONS.index(item["status"]) if item["status"] in STATUS_OPTIONS else 0,
                                format_func=lambda s: STATUS_DISPLAY[s],
                                label_visibility="collapsed",
                                key=f"tracker_sel_{item['id']}",
                            )
                            if new_status != item["status"]:
                                db.update_tracker_item(item["id"], new_status)
                                st.rerun()
                        with col_link:
                            if item.get("link"):
                                st.markdown(
                                    f'<a href="{item["link"]}" target="_blank" class="source-link" '
                                    f'style="padding:5px 10px;font-size:.78rem;margin-top:4px;display:inline-flex;">'
                                    f'↗ {item.get("link_label") or "Open"}</a>',
                                    unsafe_allow_html=True,
                                )
                        st.markdown('<hr style="margin:4px 0;border-color:var(--border-light);">', unsafe_allow_html=True)

            # ── Export checklist PDF ───────────────────────────────────────
            st.markdown("<hr>", unsafe_allow_html=True)
            st.markdown('<div class="section-label">⬇️ Export Checklist</div>', unsafe_allow_html=True)
            try:
                pdf_bytes = build_tracker_pdf(tracker_items, report, progress)
                dest_slug = report.destination_country.replace(" ", "_").lower()
                st.download_button(
                    "📕 Download checklist PDF",
                    data=pdf_bytes,
                    file_name=f"visa_checklist_{dest_slug}.pdf",
                    mime="application/pdf",
                    use_container_width=False,
                    key=f"checklist_pdf_{app_id}",
                )
            except Exception as ex:
                st.caption(f"PDF export unavailable: {ex}")

    # ── TAB 6: HISTORY ────────────────────────────────────────────────────────
    with tab6:
        _step_header("🗂", "Application History", "All saved visa applications and their reports", "done" if db.get_stats().get("total", 0) > 0 else "waiting")

        db.init_db()
        stats = db.get_stats()

        # Stats row
        sc1, sc2, sc3 = st.columns(3)
        with sc1:
            st.markdown(f'<div class="metric-card"><div class="metric-label">📁 Total Applications</div><div class="metric-value">{stats["total"]}</div></div>', unsafe_allow_html=True)
        with sc2:
            st.markdown(f'<div class="metric-card"><div class="metric-label">🇬🇧 UK Visa</div><div class="metric-value">{stats["uk"]}</div></div>', unsafe_allow_html=True)
        with sc3:
            st.markdown(f'<div class="metric-card"><div class="metric-label">🇪🇺 Schengen Visa</div><div class="metric-value">{stats["schengen"]}</div></div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        apps = db.list_applications(limit=50)

        if not apps:
            st.markdown("""
            <div class="research-card">
                <div class="rc-icon">📭</div>
                <div class="rc-title">No Applications Yet</div>
                <div class="rc-sub">Complete a visa application in Steps 1–4 and it will be saved here automatically.</div>
            </div>""", unsafe_allow_html=True)
        else:
            # Check if we're viewing a specific application
            if st.session_state.view_app_id:
                app_data = db.get_application(st.session_state.view_app_id)
                if app_data and app_data.get("report_obj"):
                    col_back, col_del = st.columns([0.85, 0.15])
                    with col_back:
                        if st.button("← Back to History List"):
                            st.session_state.view_app_id = None
                            st.rerun()
                    with col_del:
                        if st.button("🗑️ Delete", type="secondary"):
                            db.delete_application(st.session_state.view_app_id)
                            st.session_state.view_app_id = None
                            st.rerun()
                    st.markdown("<br>", unsafe_allow_html=True)
                    # Show chat history
                    if app_data.get("chat_history"):
                        with st.expander("💬 View Intake Conversation"):
                            render_chat(app_data["chat_history"])
                    render_report(app_data["report_obj"], key_prefix=f"history_{st.session_state.view_app_id}")
                else:
                    st.markdown('<div class="info-banner banner-red">❌ Could not load this application.</div>', unsafe_allow_html=True)
                    if st.button("← Back"):
                        st.session_state.view_app_id = None
                        st.rerun()
            else:
                # Application list table
                st.markdown('<div class="section-label">Saved Applications</div>', unsafe_allow_html=True)
                for app in apps:
                    from datetime import datetime
                    try:
                        dt = datetime.fromisoformat(app["created_at"]).strftime("%d %b %Y, %H:%M")
                    except Exception:
                        dt = app["created_at"][:16]
                    flag = "🇬🇧" if app["visa_target"] == "uk" else "🇪🇺"
                    dest = app["destination_country"] or "Unknown"
                    nat  = app["nationality"] or "Unknown"
                    vtype = app["visa_type"] or ("UK Visitor Visa" if app["visa_target"] == "uk" else "Schengen Visa")
                    summary_short = (app["executive_summary"] or "")[:120]
                    if len(app["executive_summary"] or "") > 120:
                        summary_short += "…"

                    st.markdown(f"""
                    <div style="background:var(--bg-raised);border:1px solid var(--border);border-radius:var(--radius-md);
                                padding:16px 20px;margin-bottom:10px;">
                        <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                            <div>
                                <span style="font-size:1.1rem;">{flag}</span>
                                <strong style="color:var(--text-primary);font-size:.95rem;margin-left:6px;">{vtype}</strong>
                                <span style="color:var(--text-muted);font-size:.8rem;margin-left:10px;">#{app['id']}</span>
                            </div>
                            <span style="color:var(--text-muted);font-size:.78rem;">{dt}</span>
                        </div>
                        <div style="margin-top:8px;display:flex;flex-wrap:wrap;gap:6px;">
                            <span class="visa-pill">🌍 {nat}</span>
                            <span class="visa-pill">✈️ {dest}</span>
                            <span class="visa-pill">🎯 {app['purpose_of_visit'] or 'N/A'}</span>
                            <span class="visa-pill" style="color:var(--success);border-color:#1a5c2a;">✅ {app['status'].title()}</span>
                        </div>
                        {f'<div style="margin-top:8px;color:var(--text-secondary);font-size:.84rem;">{summary_short}</div>' if summary_short else ''}
                    </div>
                    """, unsafe_allow_html=True)

                    if st.button(f"📋 View Report #{app['id']}", key=f"view_{app['id']}"):
                        st.session_state.view_app_id = app["id"]
                        st.rerun()


if __name__ == "__main__":
    main()
